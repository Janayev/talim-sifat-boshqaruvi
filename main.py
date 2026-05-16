import sys, os, json, shutil
from datetime import datetime, date
from pathlib import Path

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QTableWidget, QTableWidgetItem, QComboBox,
    QLineEdit, QTextEdit, QFileDialog, QMessageBox, QDialog,
    QFormLayout, QScrollArea, QFrame, QProgressBar, QTabWidget,
    QGridLayout, QSplitter, QListWidget, QListWidgetItem,
    QHeaderView, QAbstractItemView, QSizePolicy, QGroupBox,
    QDateEdit, QCheckBox, QSpinBox, QStackedWidget
)
from PyQt5.QtCore import (
    Qt, QDate, QTimer, QSize, pyqtSignal, QThread
)
from PyQt5.QtGui import (
    QFont, QColor, QPixmap, QPalette, QIcon,
    QFontDatabase, QPainter, QBrush, QPen
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io

DATA_FILE = Path.home() / ".sifat_dastur_data.json"

TASKS_RAW = [
    (1,"Sentabr","Bo'limning yillik ish rejasini ishlab chiqish va direktor tasdig'iga kiritish","Tasdiqlangan yillik ish rejasi","5-sentabr","Nizom 8(1)"),
    (2,"Sentabr","Ichki baholash ishchi guruhini shakllantirish","Direktor buyrug'i, guruh tarkibi","10-sentabr","Nizom 8(14)"),
    (3,"Sentabr","Pedagog kadrlarning tayanch ma'lumotlarini tahlil qilish va jadval tuzish","Tahliliy ma'lumotnoma","15-sentabr","Nizom 8(8)"),
    (4,"Sentabr","Pedagog kadrlarning dars beradigan fanlariga mos ixtisosligini tahlil qilish","Tahlil natijalari axboroti","20-sentabr","Nizom 8(8)"),
    (5,"Sentabr","Pedagog kadrlarning qayta tayyorlashdan o'tganligini o'rganish va monitoring jadvali","Monitoring jadvali","25-sentabr","Nizom 8(9)"),
    (6,"Sentabr","O'quv jarayonining normativ-huquqiy hujjatlarga muvofiqligini o'rganish","Tahliliy ma'lumot Ped.kengashga","30-sentabr","Nizom 8(10)"),
    (7,"Sentabr","Ta'lim sifatida xavflarni aniqlash va xavflar reyestrini shakllantirish","Xavflar reyestri","30-sentabr","Nizom 8(5)"),
    (8,"Oktabr","Ta'lim sifatini takomillashtirish bo'yicha chora-tadbirlar rejasini ishlab chiqish","Chora-tadbirlar rejasi","10-oktabr","Nizom 8(2)"),
    (9,"Oktabr","Manfaatdor tomonlar uchun ta'lim sifatiga jalb etish so'rovnomasini tashkil etish","So'rovnoma natijalari bayonnomasi","15-oktabr","Nizom 8(3)"),
    (10,"Oktabr","Pedagog kadrlar uchun zamonaviy pedagogik texnologiyalar bo'yicha o'quv-seminar","Seminar bayonnomasi","20-oktabr","Nizom 8(6)"),
    (11,"Oktabr","O'quv mashg'ulotlarining zamonaviy ped. texnologiyalar asosida tashkil etilganligini o'rganish (1-tsikl)","Kuzatuv varaqalari, tahlil","25-oktabr","Nizom 8(11)"),
    (12,"Oktabr","ARM sharoitlari va darsliklar mavjudligini tahlil qilish","ARM holati axboroti","31-oktabr","Nizom 8(12)"),
    (13,"Noyabr","Pedagog kadrlar salohiyatini oshirish bo'yicha takliflar va malaka oshirish monitoringi","Takliflar ro'yxati, monitoring","10-noyabr","Nizom 8(13)"),
    (14,"Noyabr","Xavflar bo'yicha profilaktik chora-tadbirlarni ishlab chiqish","Profilaktik tadbirlar rejasi","15-noyabr","Nizom 8(5)"),
    (15,"Noyabr","Ichki baholash ishchi guruhi bilan yig'ilish va I chorak natijalarini muhokama","Ishchi guruh bayonnomasi","20-noyabr","Nizom 8(14)"),
    (16,"Noyabr","O'quv jarayoni tahliliy ma'lumotlarini Ped. kengashga taklif kiritish","Ped. kengash taklifnomalari","25-noyabr","Nizom 8(10)"),
    (17,"Dekabr","Ped. kengashida bo'lim faoliyati bo'yicha I yarim yillik hisobot taqdim etish","Kengash bayonnomasi, hisobot","10-dekabr","Nizom 8(19)"),
    (18,"Dekabr","Davlat akkreditatsiya talablariga tayyorgarlik: ishchi guruh shakllantirish","Ishchi guruh buyrug'i","15-dekabr","Nizom 8(15)"),
    (19,"Dekabr","Pedagoglar uchun o'quv-seminar №2 (ichki nazorat asoslari)","Seminar bayonnomasi","20-dekabr","Nizom 8(6)"),
    (20,"Dekabr","Yillik ish rejasi ijrosi bo'yicha yarim yillik tahlil va tuzatishlar","Tuzatishlar kiritilgan ish rejasi","31-dekabr","Nizom 8(1)"),
    (21,"Yanvar","Davlat akkreditatsiya doirasida ta'lim jarayoni, hujjatlar, MTB va kadrlar tahlili","Tahliliy ma'lumotnoma","15-yanvar","Nizom 8(17)"),
    (22,"Yanvar","Tashqi baholash xulosalarida aniqlangan kamchiliklarni bartaraf etish rejasi","Kamchiliklarni bartaraf etish rejasi","20-yanvar","Nizom 8(18)"),
    (23,"Yanvar","O'quv mashg'ulotlarini kuzatish – II tsikl","Kuzatuv varaqalari, tahlil","25-yanvar","Nizom 8(11)"),
    (24,"Yanvar","Pedagog kadrlarning malaka oshirishga jalb etilishi monitoringi tahlili","Monitoring hisoboti","31-yanvar","Nizom 8(13)"),
    (25,"Fevral","Pedagoglar uchun o'quv-seminar №3 (innovatsion metodlar, dual ta'lim)","Seminar bayonnomasi","10-fevral","Nizom 8(6)"),
    (26,"Fevral","Ish beruvchilar bilan yig'ilish: ta'lim sifatiga oid mulohazalar","Yig'ilish bayonnomasi, takliflar","15-fevral","Nizom 8(3)"),
    (27,"Fevral","Xavflar reyestrini yangilash va profilaktik chora-tadbirlarni kuchaytirish","Yangilangan xavflar reyestri","20-fevral","Nizom 8(5)"),
    (28,"Fevral","Ped. kengashida II chorak hisobotini taqdim etish","Kengash bayonnomasi","25-fevral","Nizom 8(19)"),
    (29,"Mart","Ichki baholashni o'tkazish – I bosqich: hujjatlar tahlili va indikatorlar","Baholash jadvali","15-mart","Nizom 8(14)"),
    (30,"Mart","Davlat akkreditatsiyasiga tayyorgarlik monitoringi: talablar ijrosi tekshiruvi","Monitoring natijalar jadvali","20-mart","Nizom 8(15)"),
    (31,"Mart","Pedagoglar uchun o'quv-seminar №4 (akkreditatsiya talablari, sifat mezonlari)","Seminar bayonnomasi","25-mart","Nizom 8(6)"),
    (32,"Mart","Davlat akkreditatsiya talablarining o'z vaqtida bajarilishi bo'yicha chora-tadbirlar","Nazorat jadvali","31-mart","Nizom 8(16)"),
    (33,"Aprel","Ichki baholash – II bosqich: bo'limlar bilan yig'ilish va kamchiliklarni muhokama","Ishchi guruh bayonnomasi","10-aprel","Nizom 8(14)"),
    (34,"Aprel","Aniqlangan kamchiliklarni bartaraf etish bo'yicha takliflar ishlab chiqish","Takliflar ma'lumotnomasi","15-aprel","Nizom 8(18)"),
    (35,"Aprel","O'quv mashg'ulotlarini kuzatish – III tsikl (qiyosiy tahlil)","Qiyosiy tahlil jadvali","20-aprel","Nizom 8(11)"),
    (36,"Aprel","ARM va MTBning akkreditatsiya talablariga mosligini tekshirish","Tekshiruv dalolatnomasi","25-aprel","Nizom 8(12,17)"),
    (37,"May","Pedagoglar uchun o'quv-seminar №5 (AKT texnologiyalari)","Seminar bayonnomasi","10-may","Nizom 8(6)"),
    (38,"May","O'quvchilar bilan so'rovnoma – ta'lim sifati reytingi","So'rovnoma natijalari tahlili","15-may","Nizom 8(3)"),
    (39,"May","Ped. kengashida III chorak hisobot va ichki baholash natijalarini taqdim","Kengash bayonnomasi","25-may","Nizom 8(19)"),
    (40,"Iyun","Ichki baholashni yakunlash: yakuniy hisobot loyihasini tayyorlash","Ichki baholash hisoboti (loyiha)","10-iyun","Nizom 8(14)"),
    (41,"Iyun","Tashqi baholash xulosalaridagi kamchiliklarning bartaraf etilishi monitoringi","Monitoring axboroti","15-iyun","Nizom 8(18)"),
    (42,"Iyun","Pedagoglar uchun o'quv-seminar №6 – yillik yakunlovchi seminar","Seminar bayonnomasi","20-iyun","Nizom 8(6)"),
    (43,"Iyun","Bitiruvchilar bilan mashg'ulot: kasb tanlash va kasbiy rivojlanish","Yo'naltirish suhbati bayonnomasi","25-iyun","Nizom 8(3)"),
    (44,"Iyul","Ichki baholash hisobotini rasmiylashtirish va tasdiqlash","Tasdiqlangan ichki baholash hisoboti","10-iyul","Nizom 8(14)"),
    (45,"Iyul","Davlat akkreditatsiyasiga tayyorgarlik: barcha hujjatlar to'plamini yakunlash","Hujjatlar to'plami (papka)","20-iyul","Nizom 8(15)"),
    (46,"Iyul","Pedagog kadrlar malakasini oshirish natijalari bo'yicha yillik monitoring hisoboti","Yillik monitoring hisoboti","25-iyul","Nizom 8(13)"),
    (47,"Avgust","Kelgusi o'quv yili uchun yillik ish rejasi loyihasini ishlab chiqish","Yangi yil ish rejasi (loyiha)","10-avgust","Nizom 8(1)"),
    (48,"Avgust","Yillik ish rejasi ijrosi bo'yicha yakuniy tahlil va o'z-o'zini baholash","Yakuniy tahlil ma'lumotnomasi","15-avgust","Nizom 7-band"),
    (49,"Avgust","Ped. kengashida yillik hisobot va kelgusi yil vazifalari taklifi","Kengash bayonnomasi, yillik hisobot","25-avgust","Nizom 8(19)"),
    (50,"Avgust","Yangi o'quv yiliga tayyorgarlik: hujjatlarni arxivlash va tartiblash","Arxiv hujjatlari, yangi yil reja","31-avgust","Nizom 8(1)"),
]

QUARTERS = {
    "I chorak (Sentabr–Noyabr)": list(range(1,17)),
    "II chorak (Dekabr–Fevral)": list(range(17,29)),
    "III chorak (Mart–May)": list(range(29,40)),
    "IV chorak (Iyun–Avgust)": list(range(40,51)),
}

STATUS_LABELS = {
    "todo": "Kutilmoqda",
    "prog": "Jarayonda",
    "done": "Bajarildi",
    "late": "Muddati o'tdi",
}

STATUS_COLORS = {
    "todo":  "#888780",
    "prog":  "#BA7517",
    "done":  "#1D9E75",
    "late":  "#E24B4A",
}

NAVY = "#1B3A6B"
LIGHT_BLUE = "#E6F1FB"
GOLD = "#B8860B"

def load_data():
    if DATA_FILE.exists():
        try:
            return json.loads(DATA_FILE.read_text(encoding="utf-8"))
        except:
            pass
    return {}

def save_data(data):
    DATA_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


class HisobotWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, data, out_path, scope_ids, report_type):
        super().__init__()
        self.data = data
        self.out_path = out_path
        self.scope_ids = scope_ids
        self.report_type = report_type

    def run(self):
        try:
            path = generate_word_report(self.data, self.out_path, self.scope_ids, self.report_type)
            self.finished.emit(path)
        except Exception as e:
            self.error.emit(str(e))


def set_cell_bg(cell, hex_color):
    hex_color = hex_color.lstrip("#")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="1B3A6B"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def generate_word_report(data, out_path, scope_ids, report_type):
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    def heading(text, level=1, center=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(16 if level==1 else 14 if level==2 else 12)
        if color:
            h = color.lstrip("#")
            if len(h)==3: h = h[0]*2+h[1]*2+h[2]*2
            r,g,b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
            run.font.color.rgb = RGBColor(r,g,b)
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")
        spacing.set(qn("w:after"), "80")
        pPr.append(spacing)
        return p

    def para(text, bold=False, size=12, center=False, color=None, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if color:
            h = color.lstrip("#")
            if len(h)==3: h = h[0]*2+h[1]*2+h[2]*2
            r,g,b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
            run.font.color.rgb = RGBColor(r,g,b)
        return p

    def hline():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1B3A6B")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── MUQOVA ────────────────────────────────────────────────────────────────
    para("", center=True)
    para("BOYSUN TUMAN 1-SON TEXNIKUMI", bold=True, size=14, center=True, color=NAVY)
    para("Ta'lim sifatini ta'minlash bo'limi", bold=False, size=12, center=True, color="#555")
    hline()
    para("", center=True)

    if report_type == "haftalik":
        heading("HAFTALIK HISOBOT", center=True, color=NAVY)
        now = datetime.now()
        para(f"Hisobot sanasi: {now.strftime('%d.%m.%Y')}", center=True, size=11, italic=True, color="#555")
    elif report_type == "choraklik":
        heading("CHORAKLIK HISOBOT", center=True, color=NAVY)
        now = datetime.now()
        para(f"Hisobot sanasi: {now.strftime('%d.%m.%Y')}", center=True, size=11, italic=True, color="#555")
    else:
        heading("YILLIK HISOBOT", center=True, color=NAVY)
        now = datetime.now()
        para(f"Hisobot sanasi: {now.strftime('%d.%m.%Y')}", center=True, size=11, italic=True, color="#555")

    para("", center=True)
    para("Mas'ul shaxs: Bosh mutaxassis B. Janayev", bold=True, size=12, center=True)
    para("", center=True)
    hline()

    # ── UMUMIY STATISTIKA ─────────────────────────────────────────────────────
    doc.add_paragraph()
    heading("1. Umumiy statistika", level=2, color=NAVY)

    tasks_in_scope = [t for t in TASKS_RAW if t[0] in scope_ids]
    done_tasks = [t for t in tasks_in_scope if data.get(str(t[0]),{}).get("status") == "done"]
    prog_tasks = [t for t in tasks_in_scope if data.get(str(t[0]),{}).get("status") == "prog"]
    late_tasks = [t for t in tasks_in_scope if data.get(str(t[0]),{}).get("status") == "late"]
    todo_tasks = [t for t in tasks_in_scope if data.get(str(t[0]),{}).get("status","todo") == "todo"]

    total = len(tasks_in_scope)
    pct = round(len(done_tasks)/total*100) if total else 0

    stat_tbl = doc.add_table(rows=6, cols=2)
    stat_tbl.style = "Table Grid"
    stat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ("Ko'rsatkich", "Qiymat"),
        ("Jami chora-tadbirlar", str(total)),
        ("Bajarildi ✓", str(len(done_tasks))),
        ("Jarayonda →", str(len(prog_tasks))),
        ("Muddati o'tdi !", str(len(late_tasks))),
        (f"Umumiy ijro foizi", f"{pct}%"),
    ]
    for i, (k, v) in enumerate(rows_data):
        c0 = stat_tbl.rows[i].cells[0]
        c1 = stat_tbl.rows[i].cells[1]
        c0.text = k
        c1.text = v
        for cell in (c0, c1):
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            set_cell_borders(cell)
        if i == 0:
            set_cell_bg(c0, "1B3A6B")
            set_cell_bg(c1, "1B3A6B")
            c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            c0.paragraphs[0].runs[0].bold = True
            c1.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── CHORAKLIK PROGRESS ────────────────────────────────────────────────────
    heading("2. Choraklik progress", level=2, color=NAVY)
    q_tbl = doc.add_table(rows=len(QUARTERS)+1, cols=5)
    q_tbl.style = "Table Grid"
    headers = ["Chorak", "Jami", "Bajarildi", "Jarayonda", "Ijro %"]
    for ci, h in enumerate(headers):
        cell = q_tbl.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(cell, "1B3A6B")
        set_cell_borders(cell)

    for ri, (qname, qids) in enumerate(QUARTERS.items(), 1):
        scoped = [tid for tid in qids if tid in scope_ids]
        if not scoped:
            continue
        qd = len([t for t in scoped if data.get(str(t),{}).get("status")=="done"])
        qp = len([t for t in scoped if data.get(str(t),{}).get("status")=="prog"])
        qpct = round((qd+qp*0.5)/len(scoped)*100) if scoped else 0
        row_vals = [qname, str(len(scoped)), str(qd), str(qp), f"{qpct}%"]
        for ci, val in enumerate(row_vals):
            cell = q_tbl.rows[ri].cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_bg(cell, "EAF2FB")
            set_cell_borders(cell)

    doc.add_paragraph()

    # ── BAJARILGAN TADBIRLAR (RASMLAR BILAN) ─────────────────────────────────
    done_with_data = [t for t in tasks_in_scope if data.get(str(t[0]),{}).get("status") == "done"]
    if done_with_data:
        heading("3. Bajarilgan chora-tadbirlar", level=2, color=NAVY)

        for task in done_with_data:
            tid = task[0]
            td = data.get(str(tid), {})

            # Task header
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(f"№{tid}. {task[2]}")
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(11, 68, 124)

            # Info table
            info_tbl = doc.add_table(rows=4, cols=2)
            info_tbl.style = "Table Grid"
            info_rows = [
                ("Oy:", task[1]),
                ("Kutilgan natija:", task[3]),
                ("Muddat:", task[4]),
                ("Me'yoriy asos:", task[5]),
            ]
            for ri2, (lbl, val) in enumerate(info_rows):
                c0 = info_tbl.rows[ri2].cells[0]
                c1 = info_tbl.rows[ri2].cells[1]
                c0.text = lbl
                c1.text = val
                for cell in (c0, c1):
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    set_cell_borders(cell, "AAAAAA")
                c0.paragraphs[0].runs[0].bold = True
                set_cell_bg(c0, "F0F4FA")

            # Izoh
            note = td.get("note", "").strip()
            if note:
                p2 = doc.add_paragraph()
                run2 = p2.add_run("Izoh: ")
                run2.bold = True
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(11)
                run3 = p2.add_run(note)
                run3.font.name = "Times New Roman"
                run3.font.size = Pt(11)

            # Rasmlar
            images = td.get("images", [])
            if images:
                p3 = doc.add_paragraph()
                run4 = p3.add_run("Tasdiqlovchi rasmlar:")
                run4.bold = True
                run4.font.name = "Times New Roman"
                run4.font.size = Pt(11)

                for img_path in images:
                    if os.path.exists(img_path):
                        try:
                            img = Image.open(img_path)
                            # resize to max 14cm wide
                            max_w = 1400
                            if img.width > max_w:
                                ratio = max_w / img.width
                                img = img.resize((max_w, int(img.height*ratio)), Image.LANCZOS)
                            buf = io.BytesIO()
                            img.save(buf, format="PNG")
                            buf.seek(0)
                            doc.add_picture(buf, width=Cm(14))
                            cap = doc.add_paragraph(os.path.basename(img_path))
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.runs[0].font.size = Pt(9)
                            cap.runs[0].italic = True
                            cap.runs[0].font.color.rgb = RGBColor(100,100,100)
                        except Exception as e:
                            para(f"[Rasm yuklanmadi: {img_path}]", italic=True, color="#999")

            # Separator
            hline()

    # ── JARAYONDAGI TADBIRLAR ─────────────────────────────────────────────────
    if prog_tasks:
        heading("4. Jarayondagi tadbirlar", level=2, color=NAVY)
        p_tbl = doc.add_table(rows=len(prog_tasks)+1, cols=4)
        p_tbl.style = "Table Grid"
        for ci, h in enumerate(["№","Oy","Chora-tadbir","Muddat"]):
            cell = p_tbl.rows[0].cells[ci]
            cell.text = h
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            set_cell_bg(cell, "BA7517")
            set_cell_borders(cell)
        for ri, task in enumerate(prog_tasks, 1):
            for ci, val in enumerate([str(task[0]), task[1], task[2], task[4]]):
                cell = p_tbl.rows[ri].cells[ci]
                cell.text = val
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if ri%2==0: set_cell_bg(cell, "FFF8E7")
                set_cell_borders(cell, "BA7517")
        doc.add_paragraph()

    # ── MUDDATI O'TGAN ────────────────────────────────────────────────────────
    if late_tasks:
        heading("5. Muddati o'tgan tadbirlar", level=2, color="#A32D2D")
        l_tbl = doc.add_table(rows=len(late_tasks)+1, cols=4)
        l_tbl.style = "Table Grid"
        for ci, h in enumerate(["№","Oy","Chora-tadbir","Muddat"]):
            cell = l_tbl.rows[0].cells[ci]
            cell.text = h
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            set_cell_bg(cell, "A32D2D")
            set_cell_borders(cell)
        for ri, task in enumerate(late_tasks, 1):
            for ci, val in enumerate([str(task[0]), task[1], task[2], task[4]]):
                cell = l_tbl.rows[ri].cells[ci]
                cell.text = val
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if ri%2==0: set_cell_bg(cell, "FCEBEB")
                set_cell_borders(cell, "E24B4A")
        doc.add_paragraph()

    # ── IMZO ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    hline()
    sign_tbl = doc.add_table(rows=3, cols=2)
    sign_tbl.style = "Table Grid"
    cells_data = [
        ("Bosh mutaxassis:", "Direktor:"),
        ("_________________ B. Janayev", "_________________ R. Xolmo'minova"),
        (f"Sana: {datetime.now().strftime('%d.%m.%Y')}", "Sana: ___________"),
    ]
    for ri, (l, r) in enumerate(cells_data):
        cl = sign_tbl.rows[ri].cells[0]
        cr = sign_tbl.rows[ri].cells[1]
        cl.text = l; cr.text = r
        for cell in (cl, cr):
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_borders(cell, "FFFFFF")
        if ri == 0:
            cl.paragraphs[0].runs[0].bold = True
            cr.paragraphs[0].runs[0].bold = True

    doc.save(out_path)
    return out_path


class TaskDetailDialog(QDialog):
    saved = pyqtSignal(int, dict)

    def __init__(self, task, task_data, parent=None):
        super().__init__(parent)
        self.task = task
        self.task_data = dict(task_data)
        self.image_paths = list(task_data.get("images", []))
        self.setWindowTitle(f"№{task[0]} — Tadbir tafsilotlari")
        self.setMinimumWidth(640)
        self.setMinimumHeight(560)
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(10)
        layout.setContentsMargins(16,16,16,16)

        # Header
        hdr = QLabel(f"№{self.task[0]}. {self.task[2]}")
        hdr.setWordWrap(True)
        hdr.setStyleSheet(f"font-weight:bold;font-size:13px;color:{NAVY};padding:8px;background:#E6F1FB;border-radius:6px")
        layout.addWidget(hdr)

        # Info grid
        grid = QGroupBox("Asosiy ma'lumot")
        glay = QFormLayout(grid)
        for lbl, val in [("Oy:", self.task[1]), ("Kutilgan natija:", self.task[3]),
                          ("Muddat:", self.task[4]), ("Me'yoriy asos:", self.task[5])]:
            v = QLabel(val)
            v.setWordWrap(True)
            glay.addRow(QLabel(f"<b>{lbl}</b>"), v)
        layout.addWidget(grid)

        # Status
        s_row = QHBoxLayout()
        s_row.addWidget(QLabel("<b>Holat:</b>"))
        self.status_cb = QComboBox()
        for k, v in STATUS_LABELS.items():
            self.status_cb.addItem(v, k)
        cur = self.task_data.get("status","todo")
        idx = list(STATUS_LABELS.keys()).index(cur) if cur in STATUS_LABELS else 0
        self.status_cb.setCurrentIndex(idx)
        s_row.addWidget(self.status_cb)
        s_row.addStretch()
        layout.addLayout(s_row)

        # Note
        n_grp = QGroupBox("Izoh / Bajarilganlik tavsifi")
        n_lay = QVBoxLayout(n_grp)
        self.note_edit = QTextEdit()
        self.note_edit.setPlainText(self.task_data.get("note",""))
        self.note_edit.setMaximumHeight(90)
        self.note_edit.setPlaceholderText("Bu yerga tadbir haqida qisqacha izoh yozing...")
        n_lay.addWidget(self.note_edit)
        layout.addWidget(n_grp)

        # Images
        img_grp = QGroupBox("Tasdiqlovchi rasmlar / hujjatlar")
        img_lay = QVBoxLayout(img_grp)
        btn_row = QHBoxLayout()
        add_btn = QPushButton("+ Rasm qo'shish")
        add_btn.setStyleSheet(f"background:{NAVY};color:white;padding:5px 12px;border-radius:4px;font-weight:bold")
        add_btn.clicked.connect(self._add_image)
        clear_btn = QPushButton("Tozalash")
        clear_btn.clicked.connect(self._clear_images)
        btn_row.addWidget(add_btn)
        btn_row.addWidget(clear_btn)
        btn_row.addStretch()
        img_lay.addLayout(btn_row)

        self.img_list = QListWidget()
        self.img_list.setMaximumHeight(120)
        self.img_list.setStyleSheet("font-size:11px")
        for p in self.image_paths:
            self.img_list.addItem(QListWidgetItem(os.path.basename(p)))
        img_lay.addWidget(self.img_list)

        self.img_preview = QLabel()
        self.img_preview.setMinimumHeight(100)
        self.img_preview.setAlignment(Qt.AlignCenter)
        self.img_preview.setStyleSheet("border:1px solid #ccc;border-radius:4px;background:#fafafa")
        self.img_preview.setText("Rasm tanlang (yuqoridagi ro'yxatdan)")
        self.img_list.currentRowChanged.connect(self._preview_image)
        img_lay.addWidget(self.img_preview)
        layout.addWidget(img_grp)

        # Buttons
        btn_box = QHBoxLayout()
        save_btn = QPushButton("💾  Saqlash")
        save_btn.setStyleSheet(f"background:{NAVY};color:white;padding:8px 20px;border-radius:5px;font-weight:bold;font-size:13px")
        save_btn.clicked.connect(self._save)
        cancel_btn = QPushButton("Bekor qilish")
        cancel_btn.clicked.connect(self.reject)
        btn_box.addStretch()
        btn_box.addWidget(cancel_btn)
        btn_box.addWidget(save_btn)
        layout.addLayout(btn_box)

    def _add_image(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self, "Rasm tanlang", "",
            "Rasmlar (*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.webp)"
        )
        for path in paths:
            if path not in self.image_paths:
                self.image_paths.append(path)
                self.img_list.addItem(QListWidgetItem(os.path.basename(path)))

    def _clear_images(self):
        self.image_paths.clear()
        self.img_list.clear()
        self.img_preview.setText("Rasmlar tozalandi")

    def _preview_image(self, row):
        if 0 <= row < len(self.image_paths):
            path = self.image_paths[row]
            pix = QPixmap(path)
            if not pix.isNull():
                scaled = pix.scaled(400, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.img_preview.setPixmap(scaled)
            else:
                self.img_preview.setText("Rasm yuklanmadi")

    def _save(self):
        self.task_data["status"] = self.status_cb.currentData()
        self.task_data["note"] = self.note_edit.toPlainText().strip()
        self.task_data["images"] = self.image_paths
        self.task_data["updated"] = datetime.now().isoformat()
        self.saved.emit(self.task[0], self.task_data)
        self.accept()


class HisobotDialog(QDialog):
    def __init__(self, data, parent=None):
        super().__init__(parent)
        self.data = data
        self.setWindowTitle("Hisobot yaratish")
        self.setMinimumWidth(480)
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(20,20,20,20)

        title = QLabel("📄  Word Hisobot Yaratish")
        title.setStyleSheet(f"font-size:15px;font-weight:bold;color:{NAVY}")
        layout.addWidget(title)

        # Report type
        type_grp = QGroupBox("Hisobot turi")
        type_lay = QVBoxLayout(type_grp)
        self.type_cb = QComboBox()
        self.type_cb.addItems(["Haftalik hisobot", "Choraklik hisobot", "Yillik hisobot"])
        type_lay.addWidget(self.type_cb)
        layout.addWidget(type_grp)

        # Quarter filter
        q_grp = QGroupBox("Qamrov (chorak)")
        q_lay = QVBoxLayout(q_grp)
        self.q_all = QCheckBox("Barcha choraklar (yillik)")
        self.q_all.setChecked(True)
        self.q_all.toggled.connect(self._toggle_q)
        q_lay.addWidget(self.q_all)
        self.q_checks = []
        for qname in QUARTERS:
            cb = QCheckBox(qname)
            cb.setEnabled(False)
            q_lay.addWidget(cb)
            self.q_checks.append(cb)
        layout.addWidget(q_grp)

        # Status filter
        s_grp = QGroupBox("Tadbirlarni qo'shish holati")
        s_lay = QVBoxLayout(s_grp)
        self.s_done = QCheckBox("Bajarilgan tadbirlar (rasmlar bilan)")
        self.s_done.setChecked(True)
        self.s_prog = QCheckBox("Jarayondagi tadbirlar")
        self.s_prog.setChecked(True)
        self.s_late = QCheckBox("Muddati o'tgan tadbirlar")
        self.s_late.setChecked(True)
        for cb in [self.s_done, self.s_prog, self.s_late]:
            s_lay.addWidget(cb)
        layout.addWidget(s_grp)

        # Output path
        out_grp = QGroupBox("Saqlash joyi")
        out_lay = QHBoxLayout(out_grp)
        self.out_edit = QLineEdit()
        default_name = f"hisobot_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        self.out_edit.setText(str(Path.home() / "Desktop" / default_name))
        browse = QPushButton("...")
        browse.setMaximumWidth(36)
        browse.clicked.connect(self._browse_out)
        out_lay.addWidget(self.out_edit)
        out_lay.addWidget(browse)
        layout.addWidget(out_grp)

        self.progress = QProgressBar()
        self.progress.setRange(0,0)
        self.progress.setVisible(False)
        layout.addWidget(self.progress)

        self.status_lbl = QLabel()
        self.status_lbl.setStyleSheet("color:#1D9E75;font-weight:bold")
        layout.addWidget(self.status_lbl)

        btn_box = QHBoxLayout()
        self.gen_btn = QPushButton("🖨  Hisobot yaratish")
        self.gen_btn.setStyleSheet(f"background:{NAVY};color:white;padding:9px 22px;border-radius:5px;font-weight:bold;font-size:13px")
        self.gen_btn.clicked.connect(self._generate)
        cancel_btn = QPushButton("Yopish")
        cancel_btn.clicked.connect(self.accept)
        btn_box.addStretch()
        btn_box.addWidget(cancel_btn)
        btn_box.addWidget(self.gen_btn)
        layout.addLayout(btn_box)

    def _toggle_q(self, checked):
        for cb in self.q_checks:
            cb.setEnabled(not checked)
            if not checked:
                cb.setChecked(False)

    def _browse_out(self):
        path, _ = QFileDialog.getSaveFileName(self, "Saqlash", self.out_edit.text(), "Word hujjat (*.docx)")
        if path:
            if not path.endswith(".docx"):
                path += ".docx"
            self.out_edit.setText(path)

    def _generate(self):
        out = self.out_edit.text().strip()
        if not out:
            QMessageBox.warning(self, "Xato", "Saqlash joyini ko'rsating!")
            return

        # Determine scope
        if self.q_all.isChecked():
            scope_ids = list(range(1,51))
        else:
            scope_ids = []
            for i, (qname, qids) in enumerate(QUARTERS.items()):
                if self.q_checks[i].isChecked():
                    scope_ids.extend(qids)

        type_map = {0:"haftalik", 1:"choraklik", 2:"yillik"}
        report_type = type_map[self.type_cb.currentIndex()]

        self.progress.setVisible(True)
        self.gen_btn.setEnabled(False)
        self.status_lbl.setText("Hisobot tayyorlanmoqda...")

        self.worker = HisobotWorker(self.data, out, scope_ids, report_type)
        self.worker.finished.connect(self._on_done)
        self.worker.error.connect(self._on_error)
        self.worker.start()

    def _on_done(self, path):
        self.progress.setVisible(False)
        self.gen_btn.setEnabled(True)
        self.status_lbl.setText(f"✓ Hisobot saqlandi: {os.path.basename(path)}")
        reply = QMessageBox.question(self, "Tayyor!", f"Hisobot yaratildi!\n\n{path}\n\nOchmoqchimisiz?",
                                     QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            import subprocess, platform
            if platform.system() == "Windows":
                os.startfile(path)
            elif platform.system() == "Darwin":
                subprocess.run(["open", path])
            else:
                subprocess.run(["xdg-open", path])

    def _on_error(self, err):
        self.progress.setVisible(False)
        self.gen_btn.setEnabled(True)
        QMessageBox.critical(self, "Xato", f"Hisobot yaratishda xato:\n{err}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Ta'lim sifatini ta'minlash — Boshqaruv tizimi")
        self.setMinimumSize(1100, 720)
        self.data = load_data()
        self._build_ui()
        self._refresh_table()
        self._refresh_metrics()

    def _build_ui(self):
        # Central widget
        central = QWidget()
        self.setCentralWidget(central)
        main_lay = QVBoxLayout(central)
        main_lay.setSpacing(0)
        main_lay.setContentsMargins(0,0,0,0)

        # ── TOP BAR ───────────────────────────────────────────────────────────
        top_bar = QWidget()
        top_bar.setStyleSheet(f"background:{NAVY};color:white")
        top_bar.setFixedHeight(56)
        top_lay = QHBoxLayout(top_bar)
        top_lay.setContentsMargins(16,0,16,0)

        title_lbl = QLabel("🏫  Ta'lim sifatini ta'minlash bo'limi")
        title_lbl.setStyleSheet("color:white;font-size:15px;font-weight:bold")
        sub_lbl = QLabel("Boysun tuman 1-son texnikumi · B. Janayev · 2025–2026")
        sub_lbl.setStyleSheet("color:#B5D4F4;font-size:11px")

        title_col = QVBoxLayout()
        title_col.setSpacing(1)
        title_col.addWidget(title_lbl)
        title_col.addWidget(sub_lbl)

        top_lay.addLayout(title_col)
        top_lay.addStretch()

        hisobot_btn = QPushButton("📄  Hisobot yaratish")
        hisobot_btn.setStyleSheet("background:#B8860B;color:white;padding:7px 14px;border-radius:5px;font-weight:bold;border:none")
        hisobot_btn.clicked.connect(self._open_hisobot)
        top_lay.addWidget(hisobot_btn)

        main_lay.addWidget(top_bar)

        # ── METRICS BAR ───────────────────────────────────────────────────────
        self.metrics_bar = QWidget()
        self.metrics_bar.setStyleSheet("background:#F0F4FA;border-bottom:1px solid #D6E4F0")
        self.metrics_bar.setFixedHeight(72)
        metrics_lay = QHBoxLayout(self.metrics_bar)
        metrics_lay.setContentsMargins(16,8,16,8)
        metrics_lay.setSpacing(10)

        self.metric_widgets = {}
        for key, label, color in [
            ("total","Jami tadbir","#1B3A6B"),
            ("done","Bajarildi","#085041"),
            ("prog","Jarayonda","#633806"),
            ("late","Muddati o'tdi","#791F1F"),
            ("todo","Kutilmoqda","#444441"),
            ("pct","Ijro foizi","#1B3A6B"),
        ]:
            card = QFrame()
            card.setStyleSheet("background:white;border-radius:6px;border:1px solid #D6E4F0")
            c_lay = QVBoxLayout(card)
            c_lay.setContentsMargins(10,4,10,4)
            c_lay.setSpacing(0)
            num = QLabel("—")
            num.setStyleSheet(f"font-size:20px;font-weight:bold;color:{color}")
            lbl = QLabel(label)
            lbl.setStyleSheet("font-size:10px;color:#888")
            c_lay.addWidget(num)
            c_lay.addWidget(lbl)
            self.metric_widgets[key] = num
            metrics_lay.addWidget(card)

        main_lay.addWidget(self.metrics_bar)

        # ── TABS ──────────────────────────────────────────────────────────────
        tabs = QTabWidget()
        tabs.setStyleSheet("""
            QTabBar::tab{padding:8px 18px;font-size:12px}
            QTabBar::tab:selected{font-weight:bold;color:#1B3A6B;border-bottom:2px solid #1B3A6B}
        """)

        # Tab 1: Barcha tadbirlar
        tab1 = QWidget()
        t1_lay = QVBoxLayout(tab1)
        t1_lay.setSpacing(8)
        t1_lay.setContentsMargins(12,10,12,10)

        # Controls
        ctrl = QHBoxLayout()
        self.search_le = QLineEdit()
        self.search_le.setPlaceholderText("🔍  Tadbir yoki oy bo'yicha qidirish...")
        self.search_le.textChanged.connect(self._refresh_table)
        ctrl.addWidget(self.search_le)

        self.filter_cb = QComboBox()
        self.filter_cb.addItems(["Hammasi", "Bajarildi ✓", "Jarayonda →", "Muddati o'tdi !", "Kutilmoqda"])
        self.filter_cb.currentIndexChanged.connect(self._refresh_table)
        self.filter_cb.setMinimumWidth(160)
        ctrl.addWidget(self.filter_cb)

        self.q_filter = QComboBox()
        self.q_filter.addItem("Barcha choraklar")
        for qname in QUARTERS:
            self.q_filter.addItem(qname)
        self.q_filter.currentIndexChanged.connect(self._refresh_table)
        self.q_filter.setMinimumWidth(220)
        ctrl.addWidget(self.q_filter)

        t1_lay.addLayout(ctrl)

        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(7)
        self.table.setHorizontalHeaderLabels(["№","Oy","Chora-tadbir","Natija/Hujjat","Muddat","Me'yoriy asos","Holat"])
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.table.setColumnWidth(0, 40)
        self.table.setColumnWidth(1, 75)
        self.table.setColumnWidth(4, 90)
        self.table.setColumnWidth(5, 100)
        self.table.setColumnWidth(6, 130)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setAlternatingRowColors(True)
        self.table.verticalHeader().setVisible(False)
        self.table.setStyleSheet("""
            QTableWidget{font-size:12px;gridline-color:#E0E0E0}
            QTableWidget::item{padding:4px 6px}
            QHeaderView::section{background:#1B3A6B;color:white;font-weight:bold;padding:6px;font-size:11px;border:none}
        """)
        self.table.doubleClicked.connect(self._open_detail)
        t1_lay.addWidget(self.table)

        hint = QLabel("💡  Tadbir ustiga ikki marta bosing → tafsilot, rasm va izoh qo'shing")
        hint.setStyleSheet("color:#888;font-size:11px;padding:2px 0")
        t1_lay.addWidget(hint)
        tabs.addTab(tab1, "📋  Barcha tadbirlar (50)")

        # Tab 2: Choraklar
        tab2 = QWidget()
        t2_lay = QVBoxLayout(tab2)
        t2_lay.setContentsMargins(12,10,12,10)
        self.q_panels = {}
        for qname, qids in QUARTERS.items():
            grp = QGroupBox(qname)
            grp.setStyleSheet(f"QGroupBox{{font-weight:bold;color:{NAVY};border:1px solid #B5D4F4;border-radius:6px;margin-top:8px;padding-top:6px}}QGroupBox::title{{subcontrol-origin:margin;left:8px;padding:0 4px}}")
            g_lay = QVBoxLayout(grp)
            prog_bar = QProgressBar()
            prog_bar.setRange(0,100)
            prog_bar.setStyleSheet("QProgressBar{height:12px;border-radius:6px;background:#E6F1FB}QProgressBar::chunk{background:#1B3A6B;border-radius:6px}")
            g_lay.addWidget(prog_bar)
            stat_lbl = QLabel()
            stat_lbl.setStyleSheet("font-size:11px;color:#555")
            g_lay.addWidget(stat_lbl)
            self.q_panels[qname] = (prog_bar, stat_lbl)
            t2_lay.addWidget(grp)
        t2_lay.addStretch()
        tabs.addTab(tab2, "📊  Choraklik progress")

        main_lay.addWidget(tabs)
        self.tabs = tabs

    def _refresh_table(self):
        q = self.search_le.text().lower()
        f_idx = self.filter_cb.currentIndex()
        f_map = {0:None, 1:"done", 2:"prog", 3:"late", 4:"todo"}
        f_status = f_map[f_idx]

        q_idx = self.q_filter.currentIndex()
        if q_idx == 0:
            scope = set(range(1,51))
        else:
            qname = list(QUARTERS.keys())[q_idx-1]
            scope = set(QUARTERS[qname])

        rows = []
        for t in TASKS_RAW:
            if t[0] not in scope:
                continue
            status = self.data.get(str(t[0]),{}).get("status","todo")
            if f_status and status != f_status:
                continue
            if q and q not in (t[1]+t[2]+t[3]+t[4]+t[5]).lower():
                continue
            rows.append((t, status))

        self.table.setRowCount(len(rows))
        for ri, (t, status) in enumerate(rows):
            items = [
                QTableWidgetItem(str(t[0])),
                QTableWidgetItem(t[1]),
                QTableWidgetItem(t[2]),
                QTableWidgetItem(t[3]),
                QTableWidgetItem(t[4]),
                QTableWidgetItem(t[5]),
                QTableWidgetItem(STATUS_LABELS.get(status, status)),
            ]
            color_hex = STATUS_COLORS.get(status,"#888")
            r,g,b = int(color_hex[1:3],16), int(color_hex[3:5],16), int(color_hex[5:7],16)
            for ci, item in enumerate(items):
                item.setTextAlignment(Qt.AlignVCenter | Qt.AlignLeft)
                if ci == 6:
                    item.setForeground(QColor(r,g,b))
                    item.setFont(QFont("", -1, QFont.Bold))
                self.table.setItem(ri, ci, item)
            self.table.setRowHeight(ri, 32)

        self._refresh_metrics()
        self._refresh_quarters()

    def _refresh_metrics(self):
        statuses = [self.data.get(str(t[0]),{}).get("status","todo") for t in TASKS_RAW]
        done = statuses.count("done")
        prog = statuses.count("prog")
        late = statuses.count("late")
        todo = 50 - done - prog - late
        pct = round(done/50*100)
        self.metric_widgets["total"].setText("50")
        self.metric_widgets["done"].setText(str(done))
        self.metric_widgets["prog"].setText(str(prog))
        self.metric_widgets["late"].setText(str(late))
        self.metric_widgets["todo"].setText(str(todo))
        self.metric_widgets["pct"].setText(f"{pct}%")

    def _refresh_quarters(self):
        for qname, qids in QUARTERS.items():
            statuses = [self.data.get(str(tid),{}).get("status","todo") for tid in qids]
            done = statuses.count("done")
            prog = statuses.count("prog")
            pct = round((done + prog*0.5)/len(qids)*100)
            pb, lbl = self.q_panels[qname]
            pb.setValue(pct)
            lbl.setText(f"Bajarildi: {done}/{len(qids)}  |  Jarayonda: {prog}  |  Ijro: {pct}%")

    def _open_detail(self, index):
        row = index.row()
        item = self.table.item(row, 0)
        if not item:
            return
        tid = int(item.text())
        task = next(t for t in TASKS_RAW if t[0]==tid)
        task_data = self.data.get(str(tid), {})
        dlg = TaskDetailDialog(task, task_data, self)
        dlg.saved.connect(self._on_task_saved)
        dlg.exec_()

    def _on_task_saved(self, tid, td):
        self.data[str(tid)] = td
        save_data(self.data)
        self._refresh_table()

    def _open_hisobot(self):
        dlg = HisobotDialog(self.data, self)
        dlg.exec_()


def main():
    app = QApplication(sys.argv)
    app.setApplicationName("Ta'lim sifati boshqaruv tizimi")
    app.setStyle("Fusion")

    palette = QPalette()
    palette.setColor(QPalette.Window, QColor("#F5F7FA"))
    palette.setColor(QPalette.WindowText, QColor("#1A1A1A"))
    palette.setColor(QPalette.Base, QColor("#FFFFFF"))
    palette.setColor(QPalette.AlternateBase, QColor("#F0F4FA"))
    palette.setColor(QPalette.Highlight, QColor("#1B3A6B"))
    palette.setColor(QPalette.HighlightedText, QColor("#FFFFFF"))
    app.setPalette(palette)

    win = MainWindow()
    win.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
